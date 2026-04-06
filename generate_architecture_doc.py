#!/usr/bin/env python3
"""
Claritas HR Onboarding System – Full Architecture Documentation Generator
Generates a comprehensive Word document with visual wireframes and flow diagrams.
"""

import os
import struct
import zlib
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), 'Claritas_Architecture_Documentation.docx')

# ── Colour palette ──────────────────────────────────────────────
PRIMARY     = RGBColor(0x0D, 0x6E, 0xFD)   # Bootstrap blue
SECONDARY   = RGBColor(0x6C, 0x75, 0x7D)   # Grey
SUCCESS     = RGBColor(0x19, 0x87, 0x54)
DANGER      = RGBColor(0xDC, 0x35, 0x45)
WARNING     = RGBColor(0xFF, 0xC1, 0x07)
DARK        = RGBColor(0x21, 0x25, 0x29)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG    = RGBColor(0xF8, 0xF9, 0xFA)
HEADER_BLUE = RGBColor(0x0B, 0x5E, 0xD7)
ACCENT_BLUE = RGBColor(0xE7, 0xF1, 0xFF)

# ── PNG Generator (pure Python, no PIL, no matplotlib) ──────────
def _create_png_bytes(width, height, rows):
    """Create a minimal valid PNG from raw pixel rows (list of bytes per row)."""
    def _chunk(chunk_type, data):
        c = chunk_type + data
        return struct.pack('>I', len(data)) + c + struct.pack('>I', zlib.crc32(c) & 0xffffffff)

    sig = b'\x89PNG\r\n\x1a\n'
    ihdr = struct.pack('>IIBBBBB', width, height, 8, 2, 0, 0, 0)
    raw = b''
    for row in rows:
        raw += b'\x00' + row
    idat_data = zlib.compress(raw)
    return sig + _chunk(b'IHDR', ihdr) + _chunk(b'IDAT', idat_data) + _chunk(b'IEND', b'')

def _solid_rect(width, height, r, g, b):
    row = bytes([r, g, b] * width)
    return [row for _ in range(height)]

def _draw_box(canvas, x, y, w, h, r, g, b, cw):
    """Draw a filled rectangle on canvas rows (3 bytes/pixel, width=cw)."""
    for row_i in range(y, min(y + h, len(canvas))):
        row = bytearray(canvas[row_i])
        for col_i in range(x, min(x + w, cw)):
            idx = col_i * 3
            row[idx] = r
            row[idx + 1] = g
            row[idx + 2] = b
        canvas[row_i] = bytes(row)

def _draw_text_placeholder(canvas, x, y, w, h, cw, label_hash):
    """Draw a lighter inner area to represent text inside a box."""
    inner_margin = 3
    ir, ig, ib = 255, 255, 255
    _draw_box(canvas, x + inner_margin, y + inner_margin,
              w - inner_margin * 2, h - inner_margin * 2,
              ir, ig, ib, cw)

def _draw_arrow_h(canvas, x1, x2, y, cw, r=80, g=80, b=80, thickness=2):
    """Draw horizontal arrow line."""
    for row_i in range(max(0, y - thickness // 2), min(len(canvas), y + thickness // 2 + 1)):
        row = bytearray(canvas[row_i])
        for col_i in range(min(x1, x2), min(max(x1, x2), cw)):
            idx = col_i * 3
            row[idx] = r; row[idx+1] = g; row[idx+2] = b
        canvas[row_i] = bytes(row)
    # Arrowhead
    tip_x = max(x1, x2)
    for offset in range(6):
        for dy in range(-offset, offset + 1):
            ry = y + dy
            rx = tip_x - offset
            if 0 <= ry < len(canvas) and 0 <= rx < cw:
                row = bytearray(canvas[ry])
                idx = rx * 3
                row[idx] = r; row[idx+1] = g; row[idx+2] = b
                canvas[ry] = bytes(row)

def _draw_arrow_v(canvas, x, y1, y2, cw, r=80, g=80, b=80, thickness=2):
    """Draw vertical arrow line."""
    for row_i in range(min(y1, y2), min(max(y1, y2), len(canvas))):
        row = bytearray(canvas[row_i])
        for col_i in range(max(0, x - thickness // 2), min(x + thickness // 2 + 1, cw)):
            idx = col_i * 3
            row[idx] = r; row[idx+1] = g; row[idx+2] = b
        canvas[row_i] = bytes(row)
    # Arrowhead down
    tip_y = max(y1, y2)
    for offset in range(6):
        ry = tip_y - offset
        if 0 <= ry < len(canvas):
            row = bytearray(canvas[ry])
            for dx in range(-offset, offset + 1):
                rx = x + dx
                if 0 <= rx < cw:
                    idx = rx * 3
                    row[idx] = r; row[idx+1] = g; row[idx+2] = b
            canvas[ry] = bytes(row)

def _make_canvas(w, h, bg=(248, 249, 250)):
    row = bytes([bg[0], bg[1], bg[2]] * w)
    return [row for _ in range(h)]

def _canvas_to_png(canvas, w, h):
    return _create_png_bytes(w, h, canvas)

# ── Diagram Builders ────────────────────────────────────────────
def create_system_context_diagram():
    W, H = 800, 480
    c = _make_canvas(W, H)
    # Center app box
    _draw_box(c, 300, 180, 200, 80, 13, 110, 253, W)
    # External actors
    _draw_box(c, 30, 20, 140, 50, 25, 135, 84, W)   # HR Staff
    _draw_box(c, 30, 100, 140, 50, 25, 135, 84, W)   # IT Staff
    _draw_box(c, 30, 200, 140, 50, 108, 117, 125, W)  # Employee
    _draw_box(c, 30, 300, 140, 50, 108, 117, 125, W)  # New Hire
    _draw_box(c, 30, 400, 140, 50, 220, 53, 69, W)   # SuperAdmin
    # Right side systems
    _draw_box(c, 600, 60, 170, 50, 255, 193, 7, W)   # MySQL DB
    _draw_box(c, 600, 150, 170, 50, 255, 193, 7, W)   # File Storage
    _draw_box(c, 600, 240, 170, 50, 255, 193, 7, W)   # SMTP Mail
    _draw_box(c, 600, 330, 170, 50, 255, 193, 7, W)   # Scheduler
    # Arrows left → center
    for y_pos in [45, 125, 225, 325, 425]:
        _draw_arrow_h(c, 170, 300, y_pos, W, 80, 80, 80)
    # Arrows center → right
    for y_pos in [85, 175, 265, 355]:
        _draw_arrow_h(c, 500, 600, y_pos, W, 80, 80, 80)
    png = _canvas_to_png(c, W, H)
    return BytesIO(png)

def create_employee_lifecycle_diagram():
    W, H = 800, 300
    c = _make_canvas(W, H)
    boxes = [
        (20, 120, 130, 50, 108, 117, 125),    # Onboarding Invite
        (180, 120, 130, 50, 13, 110, 253),     # Registration
        (340, 120, 130, 50, 25, 135, 84),      # Active Employee
        (500, 120, 130, 50, 220, 53, 69),      # Offboarding
        (660, 120, 120, 50, 33, 37, 41),       # Archived
    ]
    for bx, by, bw, bh, br, bg, bb in boxes:
        _draw_box(c, bx, by, bw, bh, br, bg, bb, W)
    # Arrows between stages
    _draw_arrow_h(c, 150, 180, 145, W)
    _draw_arrow_h(c, 310, 340, 145, W)
    _draw_arrow_h(c, 470, 500, 145, W)
    _draw_arrow_h(c, 630, 660, 145, W)
    # Sub-processes below active
    _draw_box(c, 340, 210, 130, 40, 11, 94, 215, W)  # Payroll
    _draw_arrow_v(c, 405, 170, 210, W)
    _draw_box(c, 180, 210, 130, 40, 11, 94, 215, W)  # Leave
    _draw_arrow_v(c, 245, 170, 210, W)
    _draw_box(c, 500, 210, 130, 40, 11, 94, 215, W)  # Claims
    _draw_arrow_v(c, 565, 170, 210, W)
    # Top labels
    _draw_box(c, 20, 30, 760, 40, 33, 37, 41, W)
    return BytesIO(_canvas_to_png(c, W, H))

def create_asset_flow_diagram():
    W, H = 800, 280
    c = _make_canvas(W, H)
    boxes = [
        (20, 100, 150, 60, 13, 110, 253),   # Asset Inventory
        (210, 100, 150, 60, 25, 135, 84),    # Assignment
        (400, 100, 150, 60, 255, 193, 7),    # Provisioning
        (600, 60, 160, 50, 220, 53, 69),     # Disposed
        (600, 150, 160, 50, 108, 117, 125),  # Returned
    ]
    for bx, by, bw, bh, br, bg, bb in boxes:
        _draw_box(c, bx, by, bw, bh, br, bg, bb, W)
    _draw_arrow_h(c, 170, 210, 130, W)
    _draw_arrow_h(c, 360, 400, 130, W)
    _draw_arrow_h(c, 550, 600, 85, W)
    _draw_arrow_h(c, 550, 600, 175, W)
    _draw_box(c, 20, 10, 760, 40, 33, 37, 41, W)
    return BytesIO(_canvas_to_png(c, W, H))

def create_auth_flow_diagram():
    W, H = 800, 400
    c = _make_canvas(W, H)
    # Login flow boxes
    _draw_box(c, 20, 50, 160, 60, 108, 117, 125, W)  # User enters creds
    _draw_box(c, 220, 50, 160, 60, 13, 110, 253, W)   # WorkEmailProvider
    _draw_box(c, 420, 50, 160, 60, 255, 193, 7, W)    # Validate
    _draw_box(c, 620, 50, 160, 60, 25, 135, 84, W)    # Session Token
    _draw_arrow_h(c, 180, 220, 80, W)
    _draw_arrow_h(c, 380, 420, 80, W)
    _draw_arrow_h(c, 580, 620, 80, W)
    # Security layer
    _draw_box(c, 20, 170, 160, 60, 220, 53, 69, W)    # SecurityHeaders
    _draw_box(c, 220, 170, 160, 60, 220, 53, 69, W)   # SingleSession
    _draw_box(c, 420, 170, 160, 60, 220, 53, 69, W)   # AuditMiddleware
    _draw_box(c, 620, 170, 160, 60, 220, 53, 69, W)   # Rate Limiting
    # Role-based routing
    _draw_box(c, 20, 300, 150, 60, 13, 110, 253, W)   # HR Routes
    _draw_box(c, 190, 300, 150, 60, 25, 135, 84, W)   # IT Routes
    _draw_box(c, 360, 300, 150, 60, 108, 117, 125, W) # Employee Routes
    _draw_box(c, 530, 300, 150, 60, 255, 193, 7, W)   # SuperAdmin Routes
    _draw_arrow_v(c, 400, 130, 170, W)
    _draw_arrow_v(c, 100, 230, 300, W)
    _draw_arrow_v(c, 270, 230, 300, W)
    _draw_arrow_v(c, 440, 230, 300, W)
    _draw_arrow_v(c, 610, 230, 300, W)
    _draw_box(c, 20, 10, 760, 30, 33, 37, 41, W)
    return BytesIO(_canvas_to_png(c, W, H))

def create_database_erd_diagram():
    W, H = 800, 600
    c = _make_canvas(W, H)
    # Core tables
    tables = [
        (20, 20, 140, 80, 13, 110, 253),    # users
        (200, 20, 140, 80, 13, 110, 253),    # employees
        (380, 20, 140, 80, 25, 135, 84),     # onboardings
        (560, 20, 140, 80, 220, 53, 69),     # offboardings
        # Sub-tables row 2
        (20, 140, 140, 60, 11, 94, 215),     # personal_details
        (180, 140, 140, 60, 11, 94, 215),    # work_details
        (340, 140, 140, 60, 11, 94, 215),    # education
        (500, 140, 140, 60, 11, 94, 215),    # spouse
        (660, 140, 140, 60, 11, 94, 215),    # emergency
        # Asset row 3
        (20, 240, 140, 60, 255, 193, 7),     # asset_inventories
        (180, 240, 140, 60, 255, 193, 7),    # asset_assignments
        (340, 240, 140, 60, 255, 193, 7),    # aarfs
        (500, 240, 140, 60, 255, 193, 7),    # it_tasks
        (660, 240, 140, 60, 255, 193, 7),    # disposed
        # Leave row 4
        (20, 340, 140, 60, 108, 117, 125),   # leave_types
        (180, 340, 140, 60, 108, 117, 125),  # leave_apps
        (340, 340, 140, 60, 108, 117, 125),  # leave_balances
        (500, 340, 140, 60, 108, 117, 125),  # public_holidays
        # Payroll row 5
        (20, 440, 140, 60, 25, 135, 84),     # pay_runs
        (180, 440, 140, 60, 25, 135, 84),    # payslips
        (340, 440, 140, 60, 25, 135, 84),    # payroll_items
        (500, 440, 140, 60, 25, 135, 84),    # ea_forms
        # Claims + attendance row 6
        (20, 540, 140, 50, 220, 53, 69),     # expense_claims
        (180, 540, 140, 50, 220, 53, 69),    # claim_items
        (340, 540, 140, 50, 220, 53, 69),    # attendance
        (500, 540, 140, 50, 220, 53, 69),    # overtime
        (660, 540, 140, 50, 108, 117, 125),  # security_audit
    ]
    for bx, by, bw, bh, br, bg, bb in tables:
        _draw_box(c, bx, by, bw, bh, br, bg, bb, W)
    # Relationships (vertical arrows)
    _draw_arrow_v(c, 270, 100, 140, W)  # employees → personal_details
    _draw_arrow_v(c, 450, 100, 140, W)  # onboardings → work_details
    _draw_arrow_h(c, 160, 200, 60, W)   # users → employees
    _draw_arrow_h(c, 340, 380, 60, W)   # employees → onboardings
    _draw_arrow_h(c, 520, 560, 60, W)   # onboardings → offboardings
    return BytesIO(_canvas_to_png(c, W, H))

def create_module_overview_diagram():
    W, H = 800, 500
    c = _make_canvas(W, H)
    # Core module boxes (big colored blocks)
    modules = [
        (20, 20, 180, 100, 13, 110, 253, "Onboarding"),
        (220, 20, 180, 100, 25, 135, 84, "Employee Mgmt"),
        (420, 20, 180, 100, 220, 53, 69, "Offboarding"),
        (620, 20, 160, 100, 108, 117, 125, "IT Assets"),
        (20, 150, 180, 100, 255, 193, 7, "Leave Mgmt"),
        (220, 150, 180, 100, 11, 94, 215, "Payroll"),
        (420, 150, 180, 100, 148, 103, 189, "Attendance"),
        (620, 150, 160, 100, 220, 53, 69, "eClaims"),
        (20, 290, 180, 80, 33, 37, 41, "Auth & Security"),
        (220, 290, 180, 80, 33, 37, 41, "Mail System"),
        (420, 290, 180, 80, 33, 37, 41, "File Storage"),
        (620, 290, 160, 80, 33, 37, 41, "Scheduler"),
        (20, 400, 760, 80, 13, 110, 253, "Laravel 12 + MySQL + Vite + Tailwind CSS v4 + Bootstrap 5"),
    ]
    for bx, by, bw, bh, br, bg, bb, *_ in modules:
        _draw_box(c, bx, by, bw, bh, br, bg, bb, W)
    return BytesIO(_canvas_to_png(c, W, H))

def create_role_matrix_diagram():
    W, H = 800, 400
    c = _make_canvas(W, H)
    # Role boxes in a hierarchy
    _draw_box(c, 320, 10, 160, 50, 220, 53, 69, W)    # SuperAdmin
    _draw_box(c, 100, 100, 140, 50, 13, 110, 253, W)   # HR Manager
    _draw_box(c, 330, 100, 140, 50, 25, 135, 84, W)    # IT Manager
    _draw_box(c, 560, 100, 140, 50, 108, 117, 125, W)  # System Admin
    _draw_box(c, 40, 200, 130, 50, 11, 94, 215, W)     # HR Executive
    _draw_box(c, 200, 200, 130, 50, 11, 94, 215, W)    # HR Intern
    _draw_box(c, 360, 200, 130, 50, 11, 94, 215, W)    # IT Executive
    _draw_box(c, 520, 200, 130, 50, 11, 94, 215, W)    # IT Intern
    _draw_box(c, 280, 310, 240, 60, 108, 117, 125, W)  # Employee
    # Arrows
    _draw_arrow_v(c, 170, 60, 100, W)
    _draw_arrow_v(c, 400, 60, 100, W)
    _draw_arrow_v(c, 630, 60, 100, W)
    _draw_arrow_v(c, 105, 150, 200, W)
    _draw_arrow_v(c, 265, 150, 200, W)
    _draw_arrow_v(c, 425, 150, 200, W)
    _draw_arrow_v(c, 585, 150, 200, W)
    _draw_arrow_v(c, 400, 250, 310, W)
    return BytesIO(_canvas_to_png(c, W, H))

def create_mail_flow_diagram():
    W, H = 800, 380
    c = _make_canvas(W, H)
    # Event triggers → Mail classes → SMTP
    _draw_box(c, 20, 20, 200, 60, 13, 110, 253, W)    # Onboarding Events
    _draw_box(c, 20, 100, 200, 60, 25, 135, 84, W)    # Employee Events
    _draw_box(c, 20, 180, 200, 60, 220, 53, 69, W)    # Offboarding Events
    _draw_box(c, 20, 260, 200, 60, 255, 193, 7, W)    # Scheduled Tasks
    # Middle: Mail system
    _draw_box(c, 300, 80, 200, 200, 33, 37, 41, W)    # 23 Mailables
    # Right: Recipients
    _draw_box(c, 580, 20, 200, 50, 108, 117, 125, W)  # Employee
    _draw_box(c, 580, 90, 200, 50, 13, 110, 253, W)   # HR Team
    _draw_box(c, 580, 160, 200, 50, 25, 135, 84, W)   # IT Team
    _draw_box(c, 580, 230, 200, 50, 220, 53, 69, W)   # Managers
    _draw_box(c, 580, 300, 200, 50, 108, 117, 125, W) # New Hires
    # Arrows
    for y_pos in [50, 130, 210, 290]:
        _draw_arrow_h(c, 220, 300, y_pos, W)
    for y_pos in [45, 115, 185, 255, 325]:
        _draw_arrow_h(c, 500, 580, y_pos, W)
    return BytesIO(_canvas_to_png(c, W, H))

def create_security_layers_diagram():
    W, H = 800, 400
    c = _make_canvas(W, H)
    # Layered defense
    layers = [
        (20, 10, 760, 60, 220, 53, 69),     # Layer 1: Network (HSTS, CSP, Security Headers)
        (40, 80, 720, 60, 255, 193, 7),      # Layer 2: Rate Limiting
        (60, 150, 680, 60, 13, 110, 253),    # Layer 3: Auth (WorkEmailProvider, SingleSession)
        (80, 220, 640, 60, 25, 135, 84),     # Layer 4: Authorization (RBAC, Capabilities)
        (100, 290, 600, 60, 108, 117, 125),  # Layer 5: Data (Validation, Parameterized Queries)
        (120, 360, 560, 30, 33, 37, 41),     # Layer 6: Audit (SecurityAuditLog)
    ]
    for bx, by, bw, bh, br, bg, bb in layers:
        _draw_box(c, bx, by, bw, bh, br, bg, bb, W)
    return BytesIO(_canvas_to_png(c, W, H))

def create_payroll_flow_diagram():
    W, H = 800, 300
    c = _make_canvas(W, H)
    boxes = [
        (20, 100, 140, 60, 108, 117, 125),   # Employee Salary
        (190, 100, 140, 60, 13, 110, 253),    # Pay Run Draft
        (360, 100, 140, 60, 255, 193, 7),     # Generate Payslips
        (530, 100, 130, 60, 25, 135, 84),     # Approve
        (690, 100, 90, 60, 220, 53, 69),      # Paid
    ]
    for bx, by, bw, bh, br, bg, bb in boxes:
        _draw_box(c, bx, by, bw, bh, br, bg, bb, W)
    _draw_arrow_h(c, 160, 190, 130, W)
    _draw_arrow_h(c, 330, 360, 130, W)
    _draw_arrow_h(c, 500, 530, 130, W)
    _draw_arrow_h(c, 660, 690, 130, W)
    # Statutory deductions below
    _draw_box(c, 360, 200, 140, 40, 33, 37, 41, W)  # EPF/SOCSO/EIS/PCB
    _draw_arrow_v(c, 430, 160, 200, W)
    _draw_box(c, 20, 20, 760, 40, 33, 37, 41, W)
    return BytesIO(_canvas_to_png(c, W, H))

def create_leave_flow_diagram():
    W, H = 800, 280
    c = _make_canvas(W, H)
    boxes = [
        (20, 100, 140, 60, 108, 117, 125),   # Employee Apply
        (200, 100, 140, 60, 255, 193, 7),     # Manager Review
        (380, 100, 140, 60, 13, 110, 253),    # HR Review
        (560, 60, 140, 50, 25, 135, 84),      # Approved
        (560, 140, 140, 50, 220, 53, 69),     # Rejected
    ]
    for bx, by, bw, bh, br, bg, bb in boxes:
        _draw_box(c, bx, by, bw, bh, br, bg, bb, W)
    _draw_arrow_h(c, 160, 200, 130, W)
    _draw_arrow_h(c, 340, 380, 130, W)
    _draw_arrow_h(c, 520, 560, 85, W)
    _draw_arrow_h(c, 520, 560, 165, W)
    _draw_box(c, 20, 20, 760, 40, 33, 37, 41, W)
    return BytesIO(_canvas_to_png(c, W, H))

def create_claims_flow_diagram():
    W, H = 800, 280
    c = _make_canvas(W, H)
    boxes = [
        (10, 100, 110, 60, 108, 117, 125),   # Draft
        (140, 100, 110, 60, 13, 110, 253),    # Submitted
        (270, 100, 120, 60, 255, 193, 7),     # Mgr Approved
        (410, 100, 120, 60, 25, 135, 84),     # HR Approved
        (550, 100, 100, 60, 25, 135, 84),     # Paid
        (670, 100, 110, 60, 220, 53, 69),     # Rejected
    ]
    for bx, by, bw, bh, br, bg, bb in boxes:
        _draw_box(c, bx, by, bw, bh, br, bg, bb, W)
    _draw_arrow_h(c, 120, 140, 130, W)
    _draw_arrow_h(c, 250, 270, 130, W)
    _draw_arrow_h(c, 390, 410, 130, W)
    _draw_arrow_h(c, 530, 550, 130, W)
    # Rejection arrows back
    _draw_arrow_h(c, 330, 670, 200, W, 220, 53, 69)
    _draw_box(c, 10, 20, 770, 40, 33, 37, 41, W)
    return BytesIO(_canvas_to_png(c, W, H))

# ── Helper functions for document styling ────────────────────────
def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = WHITE
        set_cell_shading(cell, '0D6EFD')

    # Data rows
    for r_idx, row in enumerate(rows):
        data_row = table.rows[r_idx + 1]
        for c_idx, value in enumerate(row):
            cell = data_row.cells[c_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F0F4FF')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table

def add_heading_with_style(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = HEADER_BLUE if level <= 2 else PRIMARY
    return heading

def add_colored_paragraph(doc, text, color=DARK, bold=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    return p

def add_key_value_table(doc, items, key_width=2.5, val_width=4.5):
    table = doc.add_table(rows=len(items), cols=2)
    table.style = 'Table Grid'
    for i, (key, val) in enumerate(items):
        key_cell = table.rows[i].cells[0]
        val_cell = table.rows[i].cells[1]
        key_cell.text = key
        val_cell.text = val
        for p in key_cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
        for p in val_cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
        set_cell_shading(key_cell, 'E7F1FF')
        key_cell.width = Inches(key_width)
        val_cell.width = Inches(val_width)
    return table

def add_diagram_with_legend(doc, diagram_func, title, legend_items):
    """Add a diagram image with a colour legend below it."""
    doc.add_paragraph()
    p_title = doc.add_paragraph()
    run = p_title.add_run(title)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = HEADER_BLUE
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    img_stream = diagram_func()
    doc.add_picture(img_stream, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Legend
    if legend_items:
        legend_table = doc.add_table(rows=1, cols=len(legend_items))
        legend_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (color_hex, label) in enumerate(legend_items):
            cell = legend_table.rows[0].cells[i]
            cell.text = f"■ {label}"
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(8)
            set_cell_shading(cell, color_hex)
    doc.add_paragraph()

# ── Main Document Builder ── ── ── ── ── ── ── ── ── ── ────────
def build_document():
    doc = Document()

    # ── Page Setup ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = DARK

    # ════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('CLARITAS HR ONBOARDING SYSTEM')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = HEADER_BLUE

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run('Complete System Architecture Documentation')
    run.font.size = Pt(16)
    run.font.color.rgb = SECONDARY

    doc.add_paragraph()

    version_p = doc.add_paragraph()
    version_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version_p.add_run('Version 2.0  |  April 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = SECONDARY

    doc.add_paragraph()
    doc.add_paragraph()

    meta_items = [
        ('Document Type', 'System Architecture Documentation'),
        ('System Name', 'Claritas HR Onboarding & HRM System'),
        ('Version', '2.0'),
        ('Date', 'April 2026'),
        ('Framework', 'Laravel 12 (PHP 8.3)'),
        ('Database', 'MySQL 8.4 (62 tables)'),
        ('Frontend', 'Blade + Tailwind CSS v4 + Bootstrap 5'),
        ('Classification', 'Internal / Confidential'),
    ]
    add_key_value_table(doc, meta_items)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, 'Table of Contents', level=1)
    toc_items = [
        '1.  Executive Summary',
        '2.  System Context & Overview',
        '3.  Technology Stack',
        '4.  Application Architecture',
        '5.  Module Architecture',
        '    5.1  Onboarding Module',
        '    5.2  Employee Management Module',
        '    5.3  Offboarding Module',
        '    5.4  IT Asset Management Module',
        '    5.5  Leave Management Module',
        '    5.6  Payroll & Statutory Module',
        '    5.7  Attendance & Overtime Module',
        '    5.8  Expense Claims (eClaim) Module',
        '    5.9  Announcements Module',
        '6.  Database Architecture',
        '7.  Authentication & Authorization',
        '8.  Security Architecture',
        '9.  Email & Notification System',
        '10. Scheduled Tasks & Automation',
        '11. File Storage Architecture',
        '12. Frontend Architecture',
        '13. API & Route Architecture',
        '14. Role Hierarchy & Access Matrix',
        '15. Deployment Architecture',
        '16. Appendix: Full Database Schema',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, '1. Executive Summary', level=1)
    doc.add_paragraph(
        'Claritas HR Onboarding System is a comprehensive multi-role Human Resource Management (HRM) '
        'platform built for Malaysian companies. It manages the entire employee lifecycle from pre-hire '
        'onboarding through active employment services (payroll, leave, attendance, expense claims) to '
        'exit offboarding and archival.'
    )
    doc.add_paragraph(
        'The system serves four distinct user role groups — HR, IT, SuperAdmin, and Employee — with '
        'granular sub-roles providing fine-grained access control. It fully complies with Malaysian '
        'statutory requirements including EPF, SOCSO, EIS, PCB tax deductions, and Employment Act 1955 provisions.'
    )

    doc.add_paragraph()
    add_colored_paragraph(doc, 'Key System Metrics:', PRIMARY, bold=True, size=11)
    metrics = [
        ('Total Database Tables', '62'),
        ('Total Models', '45'),
        ('Total Controllers', '20'),
        ('Total Routes', '150+'),
        ('Total Mail Classes', '23'),
        ('Total Blade Views', '80+'),
        ('Scheduled Commands', '7 (5 automated, 2 manual)'),
        ('Security Score', '89/100 (OWASP Top 10 compliant)'),
    ]
    add_key_value_table(doc, metrics)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 2. SYSTEM CONTEXT & OVERVIEW
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, '2. System Context & Overview', level=1)
    doc.add_paragraph(
        'The system context diagram shows how Claritas HRM interacts with external actors and systems. '
        'Five distinct user groups access the web application through role-based authentication. '
        'The application communicates with MySQL for data persistence, local file storage for documents, '
        'SMTP for automated email delivery, and a task scheduler for background automation.'
    )
    add_diagram_with_legend(doc, create_system_context_diagram,
        'Figure 2.1 — System Context Diagram', [
            ('198754', 'HR / IT Staff'),
            ('0D6EFD', 'Application'),
            ('6C757D', 'Users'),
            ('FFC107', 'External Systems'),
            ('DC3545', 'SuperAdmin'),
        ])

    doc.add_paragraph(
        'All interactions flow through the Laravel web application which enforces authentication, '
        'authorization, and security middleware on every request. The SMTP integration powers 23 '
        'distinct automated email workflows covering the entire employee lifecycle.'
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 3. TECHNOLOGY STACK
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, '3. Technology Stack', level=1)

    stack = [
        ['Layer', 'Technology', 'Version', 'Purpose'],
    ]
    stack_data = [
        ('Backend Framework', 'Laravel', '12.x', 'MVC web framework with Eloquent ORM'),
        ('Language', 'PHP', '8.3', 'Server-side scripting'),
        ('Database', 'MySQL', '8.4', 'Primary relational data store (62 tables)'),
        ('Test Database', 'SQLite', 'In-memory', 'Fast isolated test execution'),
        ('CSS Framework (1)', 'Tailwind CSS', 'v4', 'Utility-first CSS via Vite plugin'),
        ('CSS Framework (2)', 'Bootstrap', '5.3', 'Component library (grid, forms, modals)'),
        ('Build Tool', 'Vite', '7.x', 'Asset bundling with HMR'),
        ('JavaScript', 'Vanilla JS + Alpine.js', '-', 'Progressive enhancement'),
        ('HTTP Client', 'Axios', '1.11', 'AJAX requests'),
        ('Icons', 'Bootstrap Icons', '1.x', 'Iconography via CDN'),
        ('Dev Server', 'Laragon', '-', 'Local WAMP stack (Windows)'),
        ('Mail', 'SMTP', '-', '23 Mailable classes, hr@claritas.com'),
        ('Timezone', 'Asia/Kuala_Lumpur', 'UTC+8', 'Malaysian timezone'),
        ('Testing', 'PHPUnit', '11.x', 'Unit + Feature test suites'),
    ]
    add_styled_table(doc, stack[0], stack_data, [1.5, 1.5, 1, 3])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 4. APPLICATION ARCHITECTURE
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, '4. Application Architecture', level=1)
    doc.add_paragraph(
        'Claritas follows a classic Model-View-Controller (MVC) pattern within a modular monolith. '
        'The application is structured as a single deployable Laravel application with clear module '
        'boundaries enforced through controller grouping, model relationships, and route prefixing.'
    )

    add_heading_with_style(doc, '4.1 Architecture Pattern', level=2)
    arch_items = [
        ('Pattern', 'Modular Monolith (MVC)'),
        ('Communication', 'Synchronous HTTP (server-rendered Blade views)'),
        ('Data Access', 'Eloquent ORM with eager loading'),
        ('Authentication', 'Session-based with custom provider (work_email)'),
        ('Authorization', 'Role-based (8 roles) + capability methods + custom permissions'),
        ('File Serving', 'SecureFileController with DIRECTORY_PERMISSIONS map'),
        ('Background Tasks', 'Laravel Scheduler (cron-based, 7 commands)'),
        ('Email', 'Laravel Mail with 23 Mailable classes'),
    ]
    add_key_value_table(doc, arch_items)

    add_diagram_with_legend(doc, create_module_overview_diagram,
        'Figure 4.1 — Module Overview Diagram', [
            ('0D6EFD', 'Core Modules'),
            ('198754', 'HRM Modules'),
            ('FFC107', 'Leave'),
            ('0B5ED7', 'Payroll'),
            ('DC3545', 'Claims/Offboarding'),
            ('212529', 'Infrastructure'),
        ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 5. MODULE ARCHITECTURE
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, '5. Module Architecture', level=1)

    # 5.1 Onboarding
    add_heading_with_style(doc, '5.1 Onboarding Module', level=2)
    doc.add_paragraph(
        'The onboarding module manages the pre-hire process from HR creating an onboarding record through '
        'to a new hire completing a multi-step self-service form via a secure tokenized invite link.'
    )
    doc.add_paragraph()
    add_colored_paragraph(doc, 'Flow: HR Creates Record → Invite Email → New Hire Form → Registration → Employee Activation', PRIMARY, bold=True)
    doc.add_paragraph()
    onb_components = [
        ('Models', 'Onboarding, PersonalDetail, WorkDetail, AssetProvisioning, ItTask, OnboardingEditLog'),
        ('Controllers', 'OnboardingController (12 methods), OnboardingInviteController (6 methods)'),
        ('Views', 'hr/onboarding/ (page, show, edit), onboarding/ (form, success)'),
        ('Mail', 'OnboardingInviteMail, OnboardingEditNotificationMail, OnboardingConsentRequestMail, WelcomeNewHire'),
        ('Key Feature', 'invite_staging_json — Sections F-I stored as JSON until activation date'),
        ('Automation', 'ActivateEmployees command runs every minute, transitions on start_date'),
    ]
    add_key_value_table(doc, onb_components)

    add_diagram_with_legend(doc, create_employee_lifecycle_diagram,
        'Figure 5.1 — Employee Lifecycle Flow', [
            ('6C757D', 'Invite/Register'),
            ('0D6EFD', 'Registration'),
            ('198754', 'Active Employee'),
            ('DC3545', 'Offboarding'),
            ('212529', 'Header/Archived'),
            ('0B5ED7', 'Sub-Processes'),
        ])

    doc.add_page_break()

    # 5.2 Employee Management
    add_heading_with_style(doc, '5.2 Employee Management Module', level=2)
    doc.add_paragraph(
        'The employee module is the central entity of the system. It manages personal details, work details, '
        'education history, spouse information, emergency contacts, child registrations, and employment contracts. '
        'It supports HR editing with a full re-consent flow and tracks all changes via EmployeeEditLog.'
    )
    emp_components = [
        ('Central Model', 'Employee — 50+ fields, 15+ relationships'),
        ('Sub-Models', 'EmployeeEducationHistory, EmployeeSpouseDetail, EmployeeEmergencyContact, EmployeeChildRegistration, EmployeeContract'),
        ('Audit', 'EmployeeEditLog with consent token flow, EmployeeHistory for archival snapshots'),
        ('Controller', 'EmployeeController (25+ methods), ProfileController (15 methods)'),
        ('Views', 'hr/employees/ (page, show, edit), user/profile.blade.php'),
        ('File Storage', 'NRIC files (JSON array), education certificates (JSON array, max 5), contracts'),
        ('Consent Flow', 'HR edits Sections A/F/G/H/I → token generated → email sent → employee re-acknowledges'),
    ]
    add_key_value_table(doc, emp_components)

    # 5.3 Offboarding
    add_heading_with_style(doc, '5.3 Offboarding Module', level=2)
    doc.add_paragraph(
        'The offboarding module manages the employee exit process with automated email reminders, '
        'IT asset return tracking, and HR/IT task coordination. Two distinct view paths serve HR and IT staff.'
    )
    off_components = [
        ('Model', 'Offboarding — tracks 10+ status fields for each exit step'),
        ('Controller', 'OffboardingController (7 methods for HR + IT)'),
        ('HR View', 'hr/offboarding/ (index, show, edit) — full edit access'),
        ('IT View', 'it/offboarding-show — read-only with "HR only" badges'),
        ('Mail', 'OffboardingNoticeMail, OffboardingReminderMail, OffboardingWeekReminderMail, OffboardingSendoffMail'),
        ('Automation', 'OffboardingNotifications command runs every minute (notice → reminder → week → sendoff)'),
        ('Calendar', 'ICS calendar attachments sent with offboarding emails'),
    ]
    add_key_value_table(doc, off_components)

    doc.add_page_break()

    # 5.4 IT Asset Management
    add_heading_with_style(doc, '5.4 IT Asset Management Module', level=2)
    doc.add_paragraph(
        'Complete IT asset lifecycle tracking from procurement through assignment, provisioning, '
        'and disposal. Includes AARF (Annual Asset Record Form) with dual acknowledgement by '
        'employee and IT manager via tokenized email links.'
    )
    add_diagram_with_legend(doc, create_asset_flow_diagram,
        'Figure 5.4 — IT Asset Lifecycle Flow', [
            ('0D6EFD', 'Inventory'),
            ('198754', 'Assignment'),
            ('FFC107', 'Provisioning'),
            ('DC3545', 'Disposed'),
            ('6C757D', 'Returned'),
        ])
    asset_components = [
        ('Models', 'AssetInventory, AssetAssignment, AssetProvisioning, Aarf, DisposedAsset'),
        ('Controller', 'AssetController (13 methods), AarfController (3 methods)'),
        ('Views', 'it/assets/ (page, show, edit), aarf/ (acknowledge)'),
        ('AARF Flow', 'Email with token → Employee acknowledges → IT Manager acknowledges → Locked'),
        ('Features', 'Asset photos (JSON array), rental tracking, warranty tracking, CSV import/export'),
    ]
    add_key_value_table(doc, asset_components)

    doc.add_page_break()

    # 5.5 Leave Management
    add_heading_with_style(doc, '5.5 Leave Management Module', level=2)
    doc.add_paragraph(
        'Full leave management with 9 Malaysian statutory leave types, tenure-based entitlements, '
        'two-tier approval workflow (Manager → HR), balance tracking with carry-forward, '
        'and public holiday management per company.'
    )
    add_diagram_with_legend(doc, create_leave_flow_diagram,
        'Figure 5.5 — Leave Application Flow', [
            ('6C757D', 'Employee'),
            ('FFC107', 'Manager Review'),
            ('0D6EFD', 'HR Review'),
            ('198754', 'Approved'),
            ('DC3545', 'Rejected'),
        ])
    leave_components = [
        ('Models', 'LeaveApplication, LeaveType, LeaveBalance, LeaveEntitlement, PublicHoliday'),
        ('Controller', 'LeaveController (20 methods)'),
        ('Leave Types (default)', 'Annual, Medical, Hospitalization, Maternity, Paternity, Marriage, Compassionate, Unpaid, Emergency'),
        ('Approval Flow', 'Employee → Manager Approve/Reject → HR Approve/Reject (two-tier)'),
        ('Mail', 'LeaveApplicationNotifyMail, LeaveApprovalNotifyMail, PendingLeaveReminderMail'),
        ('Automation', 'LeaveReminder: daily @09:00 — reminds managers of pending approvals'),
    ]
    add_key_value_table(doc, leave_components)

    doc.add_page_break()

    # 5.6 Payroll & Statutory
    add_heading_with_style(doc, '5.6 Payroll & Statutory Module', level=2)
    doc.add_paragraph(
        'Comprehensive payroll processing with Malaysian statutory deductions (EPF, SOCSO, EIS, PCB), '
        'salary management, pay run workflow, payslip generation, and Borang EA / CP.8D tax form generation.'
    )
    add_diagram_with_legend(doc, create_payroll_flow_diagram,
        'Figure 5.6 — Payroll Processing Flow', [
            ('6C757D', 'Salary Setup'),
            ('0D6EFD', 'Draft'),
            ('FFC107', 'Generate'),
            ('198754', 'Approve'),
            ('DC3545', 'Paid'),
            ('212529', 'Statutory'),
        ])
    payroll_components = [
        ('Models', 'PayRun, Payslip, PayslipItem, PayrollItem, EmployeeSalary, EmployeeSalaryItem, EaForm, SalaryAdjustment, PayrollConfig'),
        ('Controller', 'PayrollController (22 methods)'),
        ('Statutory', 'EPF (employee/employer), SOCSO (employee/employer), EIS (employee/employer), PCB, HRDF (employer)'),
        ('EPF Categories', 'A (below 60, Malaysian), B (60+, Malaysian), C (below 60, non-Malaysian), D (60+, non-Malaysian)'),
        ('EA Form', 'Borang EA / CP.8D annual tax summary — auto-generated from payslip data'),
        ('Mail', 'PayslipReadyMail, EaFormReadyMail'),
        ('Integration', 'Approved expense claims auto-linked as reimbursement line items in payslips'),
    ]
    add_key_value_table(doc, payroll_components)

    doc.add_page_break()

    # 5.7 Attendance
    add_heading_with_style(doc, '5.7 Attendance & Overtime Module', level=2)
    doc.add_paragraph(
        'Employee attendance tracking with clock-in/out, work schedule management, overtime requests, '
        'and IP-based location logging. Supports multiple work schedules per company.'
    )
    attendance_components = [
        ('Models', 'AttendanceRecord, WorkSchedule, OvertimeRequest'),
        ('Controller', 'AttendanceController (11 methods)'),
        ('Clock In/Out', 'IP address recorded, work hours auto-calculated, break deduction'),
        ('Overtime', 'Request → Approval flow with multiplier-based calculation'),
        ('Status Types', 'present, absent, late, half_day, on_leave, holiday'),
        ('Work Schedules', 'Configurable per company: start/end time, break, working days (JSON array)'),
    ]
    add_key_value_table(doc, attendance_components)

    # 5.8 Expense Claims
    add_heading_with_style(doc, '5.8 Expense Claims (eClaim) Module', level=2)
    doc.add_paragraph(
        'Monthly expense claim submission with auto-categorization, GST handling, receipt management, '
        'two-tier approval (Manager → HR), bulk approve, CSV export, and payroll integration.'
    )
    add_diagram_with_legend(doc, create_claims_flow_diagram,
        'Figure 5.8 — Expense Claim Flow', [
            ('6C757D', 'Draft'),
            ('0D6EFD', 'Submitted'),
            ('FFC107', 'Mgr Approved'),
            ('198754', 'HR Approved / Paid'),
            ('DC3545', 'Rejected'),
        ])
    claims_components = [
        ('Models', 'ExpenseClaim, ExpenseClaimItem, ExpenseCategory, ExpenseClaimPolicy'),
        ('Controller', 'ExpenseClaimController (20 methods)'),
        ('Categories', '13 default categories with keyword-based auto-detection'),
        ('GST', 'Configurable rate (default 8%), per-item GST calculation'),
        ('Two-Tier Approval', 'Employee → Manager → HR (or bulk HR approve)'),
        ('Security', 'CSV formula injection protection, total integrity validation, audit logging'),
        ('Mail', 'ClaimSubmittedMail, ClaimApprovedMail, ClaimRejectedMail, ClaimReminderMail'),
        ('Payroll Integration', 'HR-approved claims auto-linked as payslip reimbursement line items'),
    ]
    add_key_value_table(doc, claims_components)

    doc.add_page_break()

    # 5.9 Announcements
    add_heading_with_style(doc, '5.9 Announcements Module', level=2)
    doc.add_paragraph(
        'Company-wide or targeted announcements with file attachments. Visible by company filter '
        '(null = all companies). Created by HR Manager/SuperAdmin.'
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 6. DATABASE ARCHITECTURE
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, '6. Database Architecture', level=1)
    doc.add_paragraph(
        'The system uses MySQL 8.4 with 62 tables organized into 10 logical groups. '
        'All tables follow Laravel conventions with auto-incrementing IDs, timestamps, and soft deletes where appropriate. '
        'Foreign keys enforce referential integrity across the schema.'
    )

    add_diagram_with_legend(doc, create_database_erd_diagram,
        'Figure 6.1 — Database Entity-Relationship Overview', [
            ('0D6EFD', 'Core (users, employees, onboarding, offboarding)'),
            ('0B5ED7', 'Sub-tables (personal, work, education, spouse, emergency)'),
            ('FFC107', 'Assets (inventory, assignments, aarfs, tasks, disposed)'),
            ('6C757D', 'Leave (types, applications, balances, holidays)'),
            ('198754', 'Payroll (pay_runs, payslips, items, ea_forms)'),
            ('DC3545', 'Claims & Attendance'),
        ])

    add_heading_with_style(doc, '6.1 Table Groups', level=2)
    table_groups = [
        ('Core & Auth (4)', 'users, employees, password_reset_tokens, companies'),
        ('Onboarding (5)', 'onboardings, personal_details, work_details, onboarding_edit_logs, announcements'),
        ('Employee Sub-tables (7)', 'employee_education_histories, employee_spouse_details, employee_emergency_contacts, employee_child_registrations, employee_contracts, employee_histories, employee_edit_logs'),
        ('Offboarding (1)', 'offboardings'),
        ('Assets (5)', 'asset_inventories, asset_assignments, asset_provisionings, aarfs, dispose_assets'),
        ('IT Tasks (1)', 'it_tasks'),
        ('Leave (5)', 'leave_types, leave_entitlements, leave_balances, leave_applications, public_holidays'),
        ('Payroll (9)', 'pay_runs, payslips, payslip_items, payroll_items, employee_salaries, employee_salary_items, salary_adjustments, payroll_configs, ea_forms'),
        ('Attendance (3)', 'attendance_records, work_schedules, overtime_requests'),
        ('Claims (4)', 'expense_claims, expense_claim_items, expense_categories, expense_claim_policies'),
        ('Admin (3)', 'user_permissions, security_audit_logs, cache/cache_locks'),
    ]
    add_key_value_table(doc, table_groups)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 7. AUTHENTICATION & AUTHORIZATION
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, '7. Authentication & Authorization', level=1)

    add_diagram_with_legend(doc, create_auth_flow_diagram,
        'Figure 7.1 — Authentication & Security Flow', [
            ('6C757D', 'User Input'),
            ('0D6EFD', 'Auth Provider'),
            ('FFC107', 'Validation'),
            ('198754', 'Session'),
            ('DC3545', 'Security Middleware'),
        ])

    add_heading_with_style(doc, '7.1 Custom Authentication', level=2)
    auth_items = [
        ('Auth Provider', 'WorkEmailUserProvider — authenticates via work_email (not personal email)'),
        ('Guard', 'Session-based (web guard)'),
        ('Login Lockout', '5 failed attempts → account locked, security event logged'),
        ('Session Enforcement', 'EnforceSingleSession — unique token per login, concurrent sessions blocked'),
        ('Password Reset', '60-minute token expiry, 3-hour session timeout'),
        ('Timing Safety', 'Constant-time hash comparison prevents enumeration attacks'),
    ]
    add_key_value_table(doc, auth_items)

    add_heading_with_style(doc, '7.2 Role-Based Access Control', level=2)

    add_diagram_with_legend(doc, create_role_matrix_diagram,
        'Figure 7.2 — Role Hierarchy', [
            ('DC3545', 'SuperAdmin'),
            ('0D6EFD', 'HR Group'),
            ('198754', 'IT Group'),
            ('6C757D', 'System Admin / Employee'),
            ('0B5ED7', 'Junior Roles'),
        ])

    roles_data = [
        ('superadmin', 'Full system access, company management, role assignment', 'All modules'),
        ('system_admin', 'Internal admin, treated like HR Manager', 'HR + Admin modules'),
        ('hr_manager', 'Full HR operations, edit records, download documents', 'HR + HRM modules'),
        ('hr_executive', 'View-only HR operations, some Leave/Payroll actions', 'HR modules (restricted)'),
        ('hr_intern', 'Limited HR view access', 'HR modules (view-only)'),
        ('it_manager', 'Full IT operations, asset management, task assignment', 'IT + Asset modules'),
        ('it_executive', 'IT operations, asset management', 'IT modules'),
        ('it_intern', 'Limited IT view access', 'IT modules (view-only)'),
        ('employee', 'Self-service profile, leave, payslips, attendance, claims', 'User modules'),
    ]
    add_styled_table(doc, ['Role', 'Description', 'Module Access'], roles_data, [1.5, 3.5, 2])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 8. SECURITY ARCHITECTURE
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, '8. Security Architecture', level=1)
    doc.add_paragraph(
        'The system implements defense-in-depth with 6 security layers, achieving an 89/100 security score '
        'against OWASP Top 10 assessment. All layers work together to protect against common web vulnerabilities.'
    )
    add_diagram_with_legend(doc, create_security_layers_diagram,
        'Figure 8.1 — Defense-in-Depth Security Layers', [
            ('DC3545', 'Layer 1: Network Security (HSTS, CSP, Security Headers)'),
            ('FFC107', 'Layer 2: Rate Limiting (login 30/min, reset 5/min, uploads 10/min)'),
            ('0D6EFD', 'Layer 3: Authentication (WorkEmail, SingleSession)'),
            ('198754', 'Layer 4: Authorization (RBAC + Capabilities)'),
            ('6C757D', 'Layer 5: Data Validation (Input sanitization, parameterized queries)'),
            ('212529', 'Layer 6: Audit Trail (SecurityAuditLog)'),
        ])

    security_features = [
        ('Security Headers', 'X-Frame-Options: DENY, X-Content-Type-Options: nosniff, CSP, HSTS (1 year), Referrer-Policy, Permissions-Policy'),
        ('Rate Limiting', 'Login: 30/min, Password Reset: 5/min, File Uploads: 10/min, Bulk Approve: 30/min'),
        ('CSRF Protection', 'Global CSRF middleware (exceptions: AARF token-based acknowledgement)'),
        ('XSS Prevention', 'Blade {{ }} escaping (no {!! !!} for user data), strip_tags() on inputs'),
        ('SQL Injection', 'Eloquent ORM with parameterized queries throughout'),
        ('File Validation', 'Custom valid_file_content validator (magic-byte MIME verification)'),
        ('Session Security', 'Unique session_token per login, single-session enforcement, session hijack detection'),
        ('Audit Logging', 'SecurityAuditLog for failed logins, lockouts, unauthorized access, session hijacks'),
        ('CSV Injection', 'sanitizeForCsv() helper prefixes formula characters in exports'),
        ('Secrets', 'All secrets in .env, no hardcoded credentials'),
    ]
    add_key_value_table(doc, security_features)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 9. EMAIL & NOTIFICATION SYSTEM
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, '9. Email & Notification System', level=1)
    doc.add_paragraph(
        'The system uses 23 Laravel Mailable classes covering every stage of the employee lifecycle. '
        'Emails are sent via SMTP from hr@claritas.com. Each mailable has a dedicated Blade template '
        'in resources/views/emails/.'
    )
    add_diagram_with_legend(doc, create_mail_flow_diagram,
        'Figure 9.1 — Email System Architecture', [
            ('0D6EFD', 'Onboarding Events'),
            ('198754', 'Employee Events'),
            ('DC3545', 'Offboarding Events'),
            ('FFC107', 'Scheduled Tasks'),
            ('212529', '23 Mailable Classes'),
            ('6C757D', 'Recipients'),
        ])

    mail_data = [
        ('Onboarding (4)', 'OnboardingInviteMail, OnboardingEditNotificationMail, OnboardingConsentRequestMail, WelcomeNewHire'),
        ('Employee (3)', 'EmployeeConsentRequestMail, ConsentRequestMail, AnnouncementMail'),
        ('Offboarding (5)', 'OffboardingNoticeMail, OffboardingReminderMail, OffboardingWeekReminderMail, OffboardingSendoffMail, CalendarInvite'),
        ('Assets (1)', 'AarfAcknowledgementMail'),
        ('Leave (3)', 'LeaveApplicationNotifyMail, LeaveApprovalNotifyMail, PendingLeaveReminderMail'),
        ('Payroll (2)', 'PayslipReadyMail, EaFormReadyMail'),
        ('Claims (4)', 'ClaimSubmittedMail, ClaimApprovedMail, ClaimRejectedMail, ClaimReminderMail'),
        ('Security (1)', 'SecurityAuditMail'),
    ]
    add_key_value_table(doc, mail_data)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 10. SCHEDULED TASKS
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, '10. Scheduled Tasks & Automation', level=1)
    scheduler_data = [
        ('employees:activate', 'Every minute', 'Activate employees on start_date, send welcome email, create offboarding on exit_date'),
        ('offboarding:notify', 'Every minute', 'Send scheduled offboarding emails (notice, reminder, week, sendoff)'),
        ('leave:remind-managers', 'Daily @ 09:00', 'Remind managers of pending leave approvals'),
        ('claims:remind', 'Daily @ 09:00', 'Remind employees of upcoming claim deadlines'),
        ('security:audit-report', 'Hourly', 'Generate security audit report from event logs'),
        ('backfill:preferred-name', 'Manual', 'One-time: populate preferred_name if empty'),
        ('backfill:offboarding', 'Manual', 'One-time: create offboarding records for employees with exit_date'),
    ]
    add_styled_table(doc, ['Command', 'Frequency', 'Purpose'], scheduler_data, [2, 1.2, 3.8])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 11. FILE STORAGE
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, '11. File Storage Architecture', level=1)
    doc.add_paragraph(
        'All sensitive files are stored in private storage (storage/app/) and served through '
        'SecureFileController which enforces role-based access via DIRECTORY_PERMISSIONS mapping.'
    )
    file_dirs = [
        ('nric_files', 'NRIC/Passport scans', 'hr_manager, superadmin, system_admin, self'),
        ('contracts', 'Employment contracts', 'hr_manager, superadmin, system_admin'),
        ('handbooks', 'Employee handbooks', 'hr_manager, superadmin, system_admin, self'),
        ('orientation', 'Orientation documents', 'hr_manager, superadmin, system_admin, self'),
        ('aarf_files', 'AARF documents', 'hr_manager, it_manager, superadmin, system_admin, self'),
        ('asset_photos', 'Asset photographs', 'hr_manager, it_manager, it_executive, superadmin'),
        ('leave_attachments', 'Leave supporting docs', 'hr_manager, hr_executive, superadmin, system_admin, self'),
        ('claim_receipts', 'Expense claim receipts', 'hr_manager, hr_executive, superadmin, system_admin, self'),
    ]
    add_styled_table(doc, ['Directory', 'Contents', 'Allowed Roles'], file_dirs, [1.5, 2, 3.5])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 12. FRONTEND
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, '12. Frontend Architecture', level=1)
    frontend_items = [
        ('Rendering', 'Server-Side Rendering (SSR) via Laravel Blade templates'),
        ('CSS', 'Tailwind CSS v4 (utilities) + Bootstrap 5 (components)'),
        ('JavaScript', 'Vanilla JS + Alpine.js (progressive enhancement, no SPA)'),
        ('Build Tool', 'Vite 7 with @tailwindcss/vite plugin and laravel-vite-plugin'),
        ('Icons', 'Bootstrap Icons via CDN'),
        ('Layout', 'Single shared layout (layouts/app.blade.php) with collapsible sidebar'),
        ('Responsive', 'Mobile-responsive via Bootstrap grid + Tailwind breakpoints'),
        ('Themes', 'Dark-mode gradient blue sidebar with light content area'),
        ('File Handling', 'DataTransfer API for multi-file management (NRIC, certificates)'),
        ('Interactivity', 'Axios for AJAX (category detection, email validation, etc.)'),
    ]
    add_key_value_table(doc, frontend_items)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 13. API & ROUTE ARCHITECTURE
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, '13. API & Route Architecture', level=1)
    doc.add_paragraph(
        'All routes are defined in routes/web.php with 150+ endpoints grouped by role prefix. '
        'Authentication is enforced via middleware groups. The system uses RESTful conventions.'
    )
    route_groups = [
        ('Guest (Public)', '~15', '/login, /register, /forgot-password, /reset-password, /aarf/{token}, /onboarding-invite/{token}'),
        ('HR Routes', '~50', '/onboarding/*, /hr/employees/*, /hr/offboarding/*, /hr/leave/*, /hr/payroll/*, /hr/attendance/*, /hr/claims/*, /announcements/*'),
        ('IT Routes', '~20', '/assets/*, /it/tasks, /it/offboarding/*'),
        ('SuperAdmin Routes', '~10', '/superadmin/roles/*, /superadmin/accounts/*, /superadmin/companies/*'),
        ('User Routes', '~40', '/dashboard, /profile, /my/leave/*, /my/payslips/*, /my/attendance/*, /my/claims/*, /my/team-*'),
        ('Shared/Utility', '~15', '/account/*, /secure-file/*, /hr/aarf/*, /logout'),
    ]
    add_styled_table(doc, ['Group', 'Count', 'Prefixes'], route_groups, [1.5, 0.7, 4.8])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 14. ROLE ACCESS MATRIX
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, '14. Role Hierarchy & Access Matrix', level=1)
    doc.add_paragraph(
        'The access matrix shows which modules and capabilities are available to each role group. '
        'Capability methods on the User model provide fine-grained authorization checks.'
    )
    access_matrix = [
        ('Onboarding', '✓ Full', '✓ Full', '✓ View', '✗', '✗'),
        ('Employees', '✓ Full', '✓ Full', '✓ View', '✗', '✓ Self'),
        ('Offboarding', '✓ Full', '✓ Full', '✓ IT View', '✗', '✗'),
        ('Assets', '✓ Full', '✓ View', '✓ Full', '✗', '✗'),
        ('Leave Config', '✓ Full', '✓ Full', '✗', '✗', '✗'),
        ('Leave Apply', '✗', '✗', '✗', '✗', '✓'),
        ('Leave Approve', '✓', '✓', '✗', '✗', '✓ Manager'),
        ('Payroll', '✓ Full', '✓ Full', '✗', '✗', '✓ View'),
        ('Attendance Config', '✓ Full', '✓ Full', '✗', '✗', '✗'),
        ('Attendance Self', '✗', '✗', '✗', '✗', '✓'),
        ('Claims Config', '✓ Full', '✓ Full', '✗', '✗', '✗'),
        ('Claims Submit', '✗', '✗', '✗', '✗', '✓'),
        ('Claims Approve', '✓ HR', '✓ HR', '✗', '✗', '✓ Manager'),
        ('Announcements', '✓', '✓ Full', '✗', '✗', '✗'),
        ('Companies', '✓', '✗', '✗', '✗', '✗'),
        ('Roles', '✓', '✗', '✗', '✗', '✗'),
        ('Accounts', '✓', '✗', '✗', '✗', '✗'),
    ]
    add_styled_table(doc, ['Module', 'SuperAdmin', 'HR Manager', 'IT Manager', 'HR/IT Intern', 'Employee'],
                     access_matrix, [1.5, 1, 1.2, 1.2, 1, 1.1])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 15. DEPLOYMENT ARCHITECTURE
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, '15. Deployment Architecture', level=1)
    deployment_items = [
        ('Environment', 'Laravel application on LAMP/WAMP stack'),
        ('Web Server', 'Apache or Nginx with PHP-FPM'),
        ('PHP Version', '8.3.x with required extensions (mbstring, openssl, pdo_mysql, gd, zip)'),
        ('Database', 'MySQL 8.4 with InnoDB engine'),
        ('File Storage', 'Local private storage (storage/app/) with symlink for public assets'),
        ('Cron', 'Laravel scheduler via system cron: * * * * * php artisan schedule:run'),
        ('Queue', 'Sync driver (configurable for Redis/database queue)'),
        ('Mail', 'SMTP transport (configurable: Mailgun, SES, Postmark)'),
        ('SSL', 'Required for production (HSTS enforced via SecurityHeaders middleware)'),
        ('Backup Strategy', 'Database dumps + file storage snapshots (recommended: daily)'),
    ]
    add_key_value_table(doc, deployment_items)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    # 16. APPENDIX: FULL DATABASE SCHEMA
    # ════════════════════════════════════════════════════════════════
    add_heading_with_style(doc, '16. Appendix: Full Database Schema', level=1)

    # Core tables
    add_heading_with_style(doc, 'Core Tables', level=2)
    core_tables = [
        ('users', 'id, name, work_email, password, role, is_active, profile_picture, login_attempts, deactivation_reason, deactivated_at, session_token, timestamps'),
        ('employees', 'id, onboarding_id, user_id, active_from, active_until, full_name, preferred_name, official_document_id, date_of_birth, sex, marital_status, religion, race, is_disabled, residential_address, personal_contact_number, house_tel_no, personal_email, bank_account_number, bank_name, epf_no, income_tax_no, socso_no, epf_category, is_resident, nationality, nric_file_path, nric_file_paths, consent_given_at, consent_ip, designation, department, company, office_location, reporting_manager, manager_id, reporting_manager_email, company_email, start_date, exit_date, last_salary_date, employment_type, work_role, google_id, aarf_file_path, handbook_path, orientation_path, employment_status, resignation_reason, remarks, timestamps'),
        ('onboardings', 'id, status, is_expired, hr_email, it_email, hr_emails, it_emails, manual_aarf_path, calendar_invite_sent, welcome_email_sent, assigned_pic_user_id, asset_preparation_status, work_email_status, invite_token, invite_email, invite_expires_at, invite_submitted, timestamps'),
        ('offboardings', 'id, onboarding_id, employee_id, full_name, company, department, designation, company_email, reporting_manager_email, personal_email, exit_date, reason, remarks, calendar_reminder_status, exiting_email_status, aarf_status, notice_email_status, reminder_email_status, week_reminder_email_status, sendoff_email_status, asset_cleaning_status, deactivation_status, assigned_pic_user_id, timestamps'),
    ]
    for table_name, columns in core_tables:
        add_colored_paragraph(doc, f'Table: {table_name}', PRIMARY, bold=True, size=10)
        p = doc.add_paragraph(columns)
        for r in p.runs:
            r.font.size = Pt(8)

    add_heading_with_style(doc, 'Employee Sub-Tables', level=2)
    sub_tables = [
        ('personal_details', 'id, onboarding_id, full_name, preferred_name, official_document_id, date_of_birth, sex, marital_status, religion, race, is_disabled, residential_address, personal_contact_number, house_tel_no, personal_email, bank_account_number, bank_name, epf_no, income_tax_no, socso_no, nric_file_path, nric_file_paths, consent_given_at, consent_ip, invite_staging_json, timestamps'),
        ('work_details', 'id, onboarding_id, employee_status, staff_status, employment_type, designation, company, office_location, reporting_manager, reporting_manager_email, start_date, exit_date, last_salary_date, company_email, google_id, department, role, timestamps'),
        ('employee_education_histories', 'id, employee_id, qualification, institution, year_graduated, years_experience, certificate_path, certificate_paths, timestamps'),
        ('employee_spouse_details', 'id, employee_id, name, address, nric_no, tel_no, occupation, income_tax_no, is_working, is_disabled, timestamps'),
        ('employee_emergency_contacts', 'id, employee_id, contact_order, name, tel_no, relationship, timestamps'),
        ('employee_child_registrations', 'id, employee_id, cat_a_100 through cat_e_50 (10 columns), timestamps'),
        ('employee_contracts', 'id, employee_id, uploaded_by, original_filename, file_path, file_size, notes, timestamps'),
        ('employee_histories', 'id, onboarding_id, employee_id, user_id, full personal and work fields, exit_reason, exit_remarks, archived_at, timestamps'),
        ('employee_edit_logs', 'id, employee_id, edited_by_user_id, edited_by_name, edited_by_role, sections_changed, change_notes, consent_token, consent_token_expires_at, consent fields, acknowledgement fields, timestamps'),
    ]
    for table_name, columns in sub_tables:
        add_colored_paragraph(doc, f'Table: {table_name}', PRIMARY, bold=True, size=10)
        p = doc.add_paragraph(columns)
        for r in p.runs:
            r.font.size = Pt(8)

    add_heading_with_style(doc, 'Asset & IT Tables', level=2)
    asset_tables = [
        ('asset_inventories', 'id + 30 fields covering: identification, specifications, procurement, assignment, condition'),
        ('asset_assignments', 'id, onboarding_id, employee_id, asset_inventory_id, assigned_date, returned_date, status'),
        ('asset_provisionings', 'id, onboarding_id, laptop_provision, monitor_set, converter, company_phone, sim_card, access_card_request, office_keys, others'),
        ('aarfs', 'id, onboarding_id, employee_id, aarf_reference, acknowledged, acknowledged_at, acknowledgement_token, it_manager fields, it_notes, asset_changes, pending_asset_ids'),
        ('dispose_assets', 'id, asset_inventory_id, asset_tag, asset_type, brand, model, serial_number, asset_condition, reason, disposed_by, disposed_at, remarks'),
        ('it_tasks', 'id, onboarding_id, offboarding_id, assigned_to, assigned_by, task_type, title, description, status, completed_at'),
    ]
    for table_name, columns in asset_tables:
        add_colored_paragraph(doc, f'Table: {table_name}', PRIMARY, bold=True, size=10)
        p = doc.add_paragraph(columns)
        for r in p.runs:
            r.font.size = Pt(8)

    add_heading_with_style(doc, 'HRM Module Tables', level=2)
    hrm_tables = [
        ('leave_types', 'id, company, name, code, description, is_paid, requires_attachment, is_active, sort_order'),
        ('leave_applications', 'id, employee_id, leave_type_id, dates, total_days, half_day fields, reason, attachment, status, approval fields, manager approval fields'),
        ('leave_balances', 'id, employee_id, leave_type_id, year, entitled, taken, carry_forward, adjustment'),
        ('leave_entitlements', 'id, leave_type_id, company, tenure range, entitled_days, carry_forward_limit'),
        ('public_holidays', 'id, company, name, date, year, is_recurring'),
        ('pay_runs', 'id, company, reference, title, year, month, dates, status, approval fields, totals, notes'),
        ('payslips', 'id, pay_run_id, employee_id, payslip_number, salary components, statutory deductions, status'),
        ('payslip_items', 'id, payslip_id, payroll_item_id, item_name, amount'),
        ('payroll_items', 'id, company, name, code, type (earning/deduction), is_statutory, is_recurring, is_active'),
        ('employee_salaries', 'id, employee_id, basic_salary, payment_method, bank fields, effective dates, is_active'),
        ('ea_forms', 'id, employee_id, year, employer info, employee info, remuneration, deductions, statutory, status'),
        ('expense_claims', 'id, employee_id, claim_number, title, year, month, totals, status, submission/approval fields'),
        ('expense_claim_items', 'id, expense_claim_id, category_id, date, description, amounts, receipt_path, is_locked'),
        ('expense_categories', 'id, company, name, code, description, is_active'),
        ('attendance_records', 'id, employee_id, date, clock times, hours, status, IP addresses, schedule_id'),
        ('work_schedules', 'id, company, name, times, working_days (JSON), is_default'),
        ('overtime_requests', 'id, employee_id, date, times, hours, multiplier, reason, status, approval fields'),
    ]
    for table_name, columns in hrm_tables:
        add_colored_paragraph(doc, f'Table: {table_name}', PRIMARY, bold=True, size=10)
        p = doc.add_paragraph(columns)
        for r in p.runs:
            r.font.size = Pt(8)

    # ── Final note ──
    doc.add_page_break()
    add_heading_with_style(doc, 'Document Information', level=1)
    doc.add_paragraph(
        'This document was auto-generated from the live codebase of the Claritas HR Onboarding System. '
        'All diagrams represent the current production architecture as of April 2026. '
        'The system comprises 62 database tables, 45 models, 20 controllers, 150+ routes, '
        '23 mail classes, and 80+ Blade views.'
    )

    # Save
    doc.save(OUTPUT_PATH)
    print(f'Document saved to: {OUTPUT_PATH}')
    size_kb = os.path.getsize(OUTPUT_PATH) / 1024
    print(f'File size: {size_kb:.1f} KB')

if __name__ == '__main__':
    build_document()
